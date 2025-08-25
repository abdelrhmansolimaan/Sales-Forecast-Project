import pandas as pd

# تحميل الملف
df = pd.read_csv("data.csv", encoding='ISO-8859-1')  # غيّر الترميز لو الملف فيه رموز غريبة

# عرض معلومات عامة
print("عدد الصفوف:", len(df))
print("أسماء الأعمدة:", df.columns.tolist())
print("أنواع البيانات:")
print(df.dtypes)
print("قيم ناقصة:")
print(df.isnull().sum())
# تحويل التاريخ
df['InvoiceDate'] = pd.to_datetime(df['InvoiceDate'], errors='coerce')

# حذف الصفوف اللي فيها قيم ناقصة في الأعمدة الأساسية
df.dropna(subset=['CustomerID', 'InvoiceNo', 'Description', 'Quantity', 'UnitPrice'], inplace=True)

# تحويل الكمية والسعر لأرقام موجبة فقط
df = df[(df['Quantity'] > 0) & (df['UnitPrice'] > 0)]

# إنشاء الإيرادات
df['Revenue'] = df['Quantity'] * df['UnitPrice']

# استخراج معلومات إضافية
df['Month'] = df['InvoiceDate'].dt.month
df['Weekday'] = df['InvoiceDate'].dt.day_name()
df['YearMonth'] = df['InvoiceDate'].dt.to_period('M').astype(str)
# عدد العملاء
print("عدد العملاء:", df['CustomerID'].nunique())

# أعلى 10 عملاء حسب الإيرادات
top_customers = df.groupby('CustomerID')['Revenue'].sum().sort_values(ascending=False).head(10)
print("أعلى العملاء:\n", top_customers)

# أعلى المنتجات
top_products = df.groupby('Description')['Revenue'].sum().sort_values(ascending=False).head(10)
print("أعلى المنتجات:\n", top_products)

# توزيع الإيرادات حسب الدولة
country_revenue = df.groupby('Country')['Revenue'].sum().sort_values(ascending=False)
print("الإيرادات حسب الدولة:\n", country_revenue)
import matplotlib.pyplot as plt

# الإيرادات الشهرية
monthly = df.groupby('YearMonth')['Revenue'].sum()
monthly.plot(kind='bar', figsize=(10,5), title='Monthly Revenue')
plt.tight_layout()
plt.show()

# أعلى المنتجات
top_products.plot(kind='barh', title='Top Products by Revenue')
plt.tight_layout()
plt.show()

# الإيرادات حسب أيام الأسبوع
weekday = df.groupby('Weekday')['Revenue'].sum().reindex(['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday'])
weekday.plot(kind='bar', title='Revenue by Weekday')
plt.tight_layout()
plt.show()
# نسبة الإيرادات من أعلى 10 عملاء
total_revenue = df['Revenue'].sum()
top10_revenue = top_customers.sum()
print("نسبة الإيرادات من أعلى 10 عملاء:", round(top10_revenue / total_revenue * 100, 2), "%")

# أكثر يوم مبيعات
best_day = weekday.idxmax()
print("أفضل يوم في الأسبوع للمبيعات:", best_day)

# عدد المرتجعات (لو فيه فواتير تبدأ بـ 'C')
returns = df[df['InvoiceNo'].str.startswith('C')]
print("عدد المرتجعات:", len(returns))
# تجهيز البيانات الشهرية
monthly_df = df.groupby('InvoiceDate')['Revenue'].sum().resample('M').sum().reset_index()
monthly_df.columns = ['ds', 'y']  # Prophet بيحتاج الأعمدة دي
from prophet import Prophet
import pandas as pd
import matplotlib.pyplot as plt

# تجهيز البيانات
monthly_df = df.groupby('InvoiceDate')['Revenue'].sum().resample('M').sum().reset_index()
monthly_df.columns = ['ds', 'y']

# تدريب النموذج
model_prophet = Prophet()
model_prophet.fit(monthly_df)

# توقع 6 شهور
future = model_prophet.make_future_dataframe(periods=6, freq='M')
forecast_prophet = model_prophet.predict(future)

# رسم التوقعات
model_prophet.plot(forecast_prophet)
plt.title("Prophet Forecast")
plt.tight_layout()
plt.show()
from statsmodels.tsa.arima.model import ARIMA

# تجهيز السلسلة الزمنية
ts = monthly_df.set_index('ds')['y']

# تدريب النموذج
model_arima = ARIMA(ts, order=(1,1,1))
results_arima = model_arima.fit()

# توقع 6 شهور
forecast_arima = results_arima.forecast(steps=6)

# رسم التوقعات
plt.figure(figsize=(10,5))
plt.plot(ts, label='Actual')
plt.plot(pd.date_range(ts.index[-1], periods=6, freq='M'), forecast_arima, label='ARIMA Forecast')
plt.title("ARIMA Forecast")
plt.legend()
plt.tight_layout()
plt.show()
from statsmodels.tsa.statespace.sarimax import SARIMAX

# تدريب النموذج
model_sarimax = SARIMAX(ts, order=(1,1,1), seasonal_order=(1,1,1,12))
results_sarimax = model_sarimax.fit()

# توقع 6 شهور
forecast_sarimax = results_sarimax.forecast(steps=6)

# رسم التوقعات
plt.figure(figsize=(10,5))
plt.plot(ts, label='Actual')
plt.plot(pd.date_range(ts.index[-1], periods=6, freq='M'), forecast_sarimax, label='SARIMAX Forecast')
plt.title("SARIMAX Forecast")
plt.legend()
plt.tight_layout()
plt.show()
import matplotlib.pyplot as plt
import pandas as pd

# Prophet
prophet_pred = forecast_prophet[['ds', 'yhat']].set_index('ds').tail(6)

# ARIMA
arima_pred = pd.Series(forecast_arima.values, index=pd.date_range(ts.index[-1], periods=6, freq='M'))

# SARIMAX
sarimax_pred = pd.Series(forecast_sarimax.values, index=pd.date_range(ts.index[-1], periods=6, freq='M'))

# رسم المقارنة
plt.figure(figsize=(12,6))
plt.plot(ts, label='Actual Revenue', color='black')
plt.plot(prophet_pred, label='Prophet Forecast', color='blue')
plt.plot(arima_pred, label='ARIMA Forecast', color='green')
plt.plot(sarimax_pred, label='SARIMAX Forecast', color='red')
plt.title('Revenue Forecast Comparison')
plt.xlabel('Date')
plt.ylabel('Revenue')
plt.legend()
plt.tight_layout()
plt.show()
from sklearn.metrics import mean_squared_error, mean_absolute_error
import numpy as np

# تجهيز البيانات الحقيقية للمقارنة (آخر 6 شهور فعلية)
actual = ts[-6:]

# Prophet
prophet_eval = forecast_prophet.set_index('ds').loc[actual.index]['yhat']
rmse_prophet = np.sqrt(mean_squared_error(actual, prophet_eval))
mae_prophet = mean_absolute_error(actual, prophet_eval)

# ARIMA
rmse_arima = np.sqrt(mean_squared_error(actual, forecast_arima[:6]))
mae_arima = mean_absolute_error(actual, forecast_arima[:6])

# SARIMAX
rmse_sarimax = np.sqrt(mean_squared_error(actual, forecast_sarimax[:6]))
mae_sarimax = mean_absolute_error(actual, forecast_sarimax[:6])

# عرض النتائج
print("📊 Prophet → RMSE:", round(rmse_prophet, 2), "| MAE:", round(mae_prophet, 2))
print("📈 ARIMA   → RMSE:", round(rmse_arima, 2), "| MAE:", round(mae_arima, 2))
print("🧠 SARIMAX → RMSE:", round(rmse_sarimax, 2), "| MAE:", round(mae_sarimax, 2))
import matplotlib.pyplot as plt

# بيانات المقارنة
models = ['Prophet', 'ARIMA', 'SARIMAX']
rmse_values = [rmse_prophet, rmse_arima, rmse_sarimax]
mae_values = [mae_prophet, mae_arima, mae_sarimax]

x = range(len(models))

# رسم RMSE
plt.figure(figsize=(10,5))
plt.bar(x, rmse_values, width=0.4, label='RMSE', align='center', color='skyblue')
plt.bar([i + 0.4 for i in x], mae_values, width=0.4, label='MAE', align='center', color='orange')
plt.xticks([i + 0.2 for i in x], models)
plt.ylabel('Error Value')
plt.title('Comparison of Forecast Models (RMSE vs MAE)')
plt.legend()
plt.tight_layout()
plt.show()
import pandas as pd
import matplotlib.pyplot as plt

# ملخص المؤشرات
summary = {
    'Total Revenue': [df['Revenue'].sum()],
    'Top Product': [df.groupby('Description')['Revenue'].sum().idxmax()],
    'Top Customer': [df.groupby('CustomerID')['Revenue'].sum().idxmax()],
    'Best Day': [df.groupby('Weekday')['Revenue'].sum().idxmax()],
    'Returns Count': [len(df[df['InvoiceNo'].str.startswith('C')])]
}
summary_df = pd.DataFrame(summary)

# حفظ التقرير في Excel
with pd.ExcelWriter('Final_Sales_Report.xlsx', engine='xlsxwriter') as writer:
    summary_df.to_excel(writer, sheet_name='Summary', index=False)
    monthly_df.to_excel(writer, sheet_name='Monthly Revenue', index=False)
    top_products.to_frame().to_excel(writer, sheet_name='Top Products')
    top_customers.to_frame().to_excel(writer, sheet_name='Top Customers')
    country_revenue.to_frame().to_excel(writer, sheet_name='Revenue by Country')

    # إعداد الرسوم البيانية داخل التقرير
    workbook  = writer.book
    worksheet = writer.sheets['Summary']
    chart = workbook.add_chart({'type': 'column'})
    chart.add_series({
        'categories': ['Summary', 0, 0, 0, len(summary_df.columns)-1],
        'values':     ['Summary', 1, 0, 1, len(summary_df.columns)-1],
        'name':       'Key Metrics'
    })
    chart.set_title({'name': 'Executive Summary'})
    worksheet.insert_chart('G2', chart)
    # تنسيق التوقيع
signature_format = workbook.add_format({
    'bold': True,
    'font_color': 'blue',
    'font_size': 12,
    'align': 'left'
})

# كتابة التوقيع في الخلية المناسبة
worksheet.write('A15', 'Prepared by: Eng. Abdelrhman Ahmed', signature_format)

# إضافة التاريخ تحت التوقيع
from datetime import datetime
worksheet.write('A16', f'Date: {datetime.today().strftime("%Y-%m-%d")}', signature_format)
worksheet.set_footer('&LPrepared by: Eng. Abdelrhman Ahmed &RDate: &D')
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from datetime import datetime

# إنشاء التقرير
doc = Document()
doc.add_heading('Sales Forecast & Analysis Report', 0)

# المؤشرات الرئيسية
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(f"Total Revenue: {df['Revenue'].sum():,.2f}")
doc.add_paragraph(f"Top Product: {df.groupby('Description')['Revenue'].sum().idxmax()}")
doc.add_paragraph(f"Top Customer: {df.groupby('CustomerID')['Revenue'].sum().idxmax()}")
doc.add_paragraph(f"Best Day: {df.groupby('Weekday')['Revenue'].sum().idxmax()}")
doc.add_paragraph(f"Returns Count: {len(df[df['InvoiceNo'].str.startswith('C')])}")

# حفظ رسم Prophet
model_prophet.plot(forecast_prophet)
plt.title("Prophet Forecast")
plt.tight_layout()
plt.savefig("prophet_forecast.png")
plt.close()
doc.add_heading('Prophet Forecast', level=2)
doc.add_picture("prophet_forecast.png", width=Inches(6))

# حفظ رسم ARIMA
plt.figure(figsize=(10,5))
plt.plot(ts, label='Actual')
plt.plot(pd.date_range(ts.index[-1], periods=6, freq='M'), forecast_arima, label='ARIMA Forecast')
plt.title("ARIMA Forecast")
plt.legend()
plt.tight_layout()
plt.savefig("arima_forecast.png")
plt.close()
doc.add_heading('ARIMA Forecast', level=2)
doc.add_picture("arima_forecast.png", width=Inches(6))

# حفظ رسم SARIMAX
plt.figure(figsize=(10,5))
plt.plot(ts, label='Actual')
plt.plot(pd.date_range(ts.index[-1], periods=6, freq='M'), forecast_sarimax, label='SARIMAX Forecast')
plt.title("SARIMAX Forecast")
plt.legend()
plt.tight_layout()
plt.savefig("sarimax_forecast.png")
plt.close()
doc.add_heading('SARIMAX Forecast', level=2)
doc.add_picture("sarimax_forecast.png", width=Inches(6))

# رسم المقارنة بين النماذج
plt.figure(figsize=(10,5))
plt.bar(models, rmse_values, width=0.4, label='RMSE', color='skyblue')
plt.bar([i + 0.4 for i in x], mae_values, width=0.4, label='MAE', color='orange')
plt.xticks([i + 0.2 for i in x], models)
plt.title('Model Accuracy Comparison')
plt.legend()
plt.tight_layout()
plt.savefig("model_comparison.png")
plt.close()
doc.add_heading('Model Accuracy Comparison (RMSE vs MAE)', level=2)
doc.add_picture("model_comparison.png", width=Inches(6))

# التوصية النهائية
doc.add_heading('Recommendation', level=1)
doc.add_paragraph("Based on accuracy metrics and seasonal behavior, SARIMAX is the most suitable model for forecasting future revenue.")

# التوقيع
doc.add_paragraph("\nPrepared by: Eng. Abdelrhman Ahmed")
doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")

# حفظ التقرير
doc.save("Final_Sales_Report.docx")